"""
语言学习助手 - GUI 版
=====================
  S+D+C  → 西班牙语：打开西语助手查词 + 保存笔记
  S+C    → 西班牙语：Google翻译 + 保存笔记
  E+C    → 英语：Google翻译 + 保存笔记

  关闭窗口 / 点击退出按钮 → 安全退出（自动写入时间戳）

依赖安装：pip install keyboard pyperclip python-docx plyer requests
"""

import sys
import os
import re
import time
import platform
import threading
import configparser
import subprocess
from datetime import datetime

import tkinter as tk
from tkinter import scrolledtext, font as tkfont

import requests
import pyperclip
import keyboard
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from plyer import notification

# ============================================================
# 配置加载（自动创建 config.ini，容错读取）
# ============================================================
_CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.ini")
_DEFAULT_EXE = r"C:\Program Files\EsHelper\eshelper.exe"


def _load_config() -> str:
    """
    读取 config.ini 中的 eshelper 路径。
    文件不存在或字段缺失时自动写入默认值并返回默认路径。
    """
    cfg = configparser.ConfigParser()
    if os.path.exists(_CONFIG_PATH):
        try:
            cfg.read(_CONFIG_PATH, encoding="utf-8")
            return cfg.get("Settings", "eshelper_path")
        except (configparser.NoSectionError, configparser.NoOptionError,
                configparser.MissingSectionHeaderError):
            pass  # 文件损坏或字段缺失，重建

    # 写入默认配置
    cfg["Settings"] = {"eshelper_path": _DEFAULT_EXE}
    try:
        with open(_CONFIG_PATH, "w", encoding="utf-8") as f:
            cfg.write(f)
    except OSError:
        pass
    return _DEFAULT_EXE


ESHELPER_EXE = _load_config()

# ============================================================
# 全局状态
# ============================================================
processing_lock     = threading.Lock()
_entry_counter      = 1
_entry_counter_lock = threading.Lock()

# GUI 回调（由 start_gui 注入，供后台线程安全更新界面）
_gui_log_callback    = None   # callable(text: str)
_gui_count_callback  = None   # callable(n: int)
_gui_status_callback = None   # callable(msg: str)


def _gui_log(text: str):
    if _gui_log_callback:
        _gui_log_callback(str(text))


def _gui_set_count(n: int):
    if _gui_count_callback:
        _gui_count_callback(n)


def _gui_set_status(msg: str):
    if _gui_status_callback:
        _gui_status_callback(msg)


def next_entry_number() -> int:
    global _entry_counter
    with _entry_counter_lock:
        n = _entry_counter
        _entry_counter += 1
    return n


# ============================================================
# 词性对照表
# ============================================================
POS_MAP_ES = {
    "verb": "v.", "transitive verb": "tr.", "intransitive verb": "intr.",
    "reflexive verb": "Prnl.", "pronominal verb": "Prnl.",
    "auxiliary verb": "v.aux.", "noun": "n.", "masculine noun": "m.",
    "feminine noun": "f.", "proper noun": "n.prop.", "adjective": "adj.",
    "adverb": "adv.", "preposition": "prep.", "conjunction": "conj.",
    "pronoun": "pron.", "interjection": "interj.", "exclamation": "interj.",
    "article": "art.", "numeral": "num.", "participle": "part.",
    "past participle": "p.p.",
}

POS_MAP_EN = {
    "verb": "v.", "transitive verb": "Vt.", "intransitive verb": "Vi.",
    "auxiliary verb": "v.aux.", "modal verb": "v.mod.", "noun": "n.",
    "proper noun": "n.prop.", "adjective": "adj.", "adverb": "adv.",
    "preposition": "prep.", "conjunction": "conj.", "pronoun": "pron.",
    "interjection": "interj.", "exclamation": "interj.", "numeral": "num.",
    "article": "art.", "abbreviation": "abbr.", "plural": "pl.",
}

# ============================================================
# 工具函数：桌面路径 / 文档编号
# ============================================================
def get_desktop_path() -> str:
    if platform.system() == "Windows":
        try:
            import winreg
            key = winreg.OpenKey(
                winreg.HKEY_CURRENT_USER,
                r"Software\Microsoft\Windows\CurrentVersion\Explorer\User Shell Folders",
            )
            val, _ = winreg.QueryValueEx(key, "Desktop")
            return os.path.expandvars(val)
        except Exception:
            pass
    return os.path.expanduser("~/Desktop")


def get_new_doc_info() -> tuple:
    desktop = get_desktop_path()
    path = os.path.join(desktop, "学习笔记.docx")
    if not os.path.exists(path):
        return path, "学习笔记"
    idx = 2
    while True:
        title = f"学习笔记{idx}"
        path  = os.path.join(desktop, f"{title}.docx")
        if not os.path.exists(path):
            return path, title
        idx += 1


DOC_PATH, DOC_TITLE = get_new_doc_info()

# ============================================================
# Word 文档操作
# ============================================================
def _apply_font(run, bold=False, italic=False, size_pt=15, color=None):
    run.font.name   = "等线"
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn("w:eastAsia"), "等线")


def add_mixed_runs(paragraph, text: str, size_pt=15, italic=False):
    """拉丁字母段加粗，其余不加粗，全部等线。"""
    for seg in re.split(r"([A-Za-z]+)", text):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        _apply_font(run, bold=bool(re.fullmatch(r"[A-Za-z]+", seg)),
                    italic=italic, size_pt=size_pt)


def _add_blue_bottom_border(paragraph):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "16")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "0070C0")
    pBdr.append(bot)
    pPr.append(pBdr)


def create_doc_with_title(title: str) -> Document:
    doc = Document()
    p   = doc.add_paragraph()
    run = p.add_run(title)
    _apply_font(run, bold=True, size_pt=19, color=RGBColor(0x00, 0x70, 0xC0))
    _add_blue_bottom_border(p)
    return doc


def append_timestamp():
    if not os.path.exists(DOC_PATH):
        return
    try:
        doc = Document(DOC_PATH)
        p   = doc.add_paragraph()
        now = datetime.now().strftime("%Y/%m/%d  %H:%M")
        _apply_font(p.add_run(now), italic=True, size_pt=13)
        doc.save(DOC_PATH)
        _gui_log(f"📅 已记录结束时间：{now}\n")
    except Exception as e:
        _gui_log(f"⚠️ 写入时间戳失败：{e}\n")


def save_to_doc(text: str, def_lines: list) -> int:
    n            = next_entry_number()
    word_display = text[0].upper() + text[1:]
    full_line    = f"{n}| {word_display} \u2014\u2014 {' || '.join(def_lines)}"

    doc = Document(DOC_PATH) if os.path.exists(DOC_PATH) else create_doc_with_title(DOC_TITLE)
    add_mixed_runs(doc.add_paragraph(), full_line, size_pt=15)
    doc.save(DOC_PATH)
    return n

# ============================================================
# 翻译 API
# ============================================================
def get_google_definition(text: str, source_lang: str) -> list:
    is_word = len(text.split()) <= 4
    pos_map = POS_MAP_ES if source_lang == "es" else POS_MAP_EN
    params  = [("client","gtx"),("sl",source_lang),("tl","zh-CN"),("dt","t"),("dj","1"),("q",text)]
    if is_word:
        params.append(("dt", "bd"))

    try:
        resp = requests.get(
            "https://translate.googleapis.com/translate_a/single",
            params=params, timeout=6,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as e:
        _gui_log(f"⚠️ Google 翻译失败：{e}\n")
        return ["（释义获取失败）"]

    if isinstance(data, dict) and is_word:
        lines = [
            f"{pos_map.get(e.get('pos','').lower(), e.get('pos',''))}{'，'.join(str(t) for t in e.get('terms',[])[:5])}"
            for e in data.get("dict", []) if e.get("terms")
        ]
        if lines:
            return lines

    if isinstance(data, dict):
        plain = "".join(s.get("trans","") for s in data.get("sentences",[]))
    else:
        try:
            plain = "".join(seg[0] for seg in data[0] if seg[0])
        except Exception:
            plain = ""
    return [plain] if plain else ["（释义获取失败）"]

# ============================================================
# 西语助手
# ============================================================
def open_in_xiyuzhushou(word: str):
    if os.path.exists(ESHELPER_EXE):
        try:
            subprocess.Popen([ESHELPER_EXE, "-w", word])
            _gui_log("   📖 西语助手已启动\n")
        except Exception as e:
            _gui_log(f"   ⚠️ 启动西语助手失败：{e}\n")
    else:
        _gui_log(f"   ⚠️ 未找到西语助手：{ESHELPER_EXE}\n")

# ============================================================
# 核心处理
# ============================================================
def process(source_lang: str, open_app: bool = False):
    # ── FIX: 锁忙时给出反馈，不再静默返回 ──────────────────
    if not processing_lock.acquire(blocking=False):
        _gui_log("⏳ 上一条还在处理中，请稍候再试。\n")
        return

    try:
        _gui_set_status("正在捕获文本…")
        pyperclip.copy("")
        time.sleep(0.15)
        keyboard.send("ctrl+c")

        text = ""
        for _ in range(10):
            time.sleep(0.05)
            content = pyperclip.paste().strip()
            if content:
                text = content
                break

        if not text:
            _gui_log("⚠️ 未检测到选中文本！请先选中文字再按快捷键。\n")
            _gui_set_status("待机中")
            return

        if open_app:
            threading.Thread(target=open_in_xiyuzhushou, args=(text,), daemon=True).start()

        _gui_set_status("正在翻译…")
        def_lines    = get_google_definition(text, source_lang)
        word_display = text[0].upper() + text[1:]
        defs_str     = " || ".join(def_lines)

        _gui_set_status("正在写入文档…")
        n = save_to_doc(text, def_lines)
        _gui_set_count(n)

        # 终端日志
        _gui_log(f"\n{'─' * 50}\n")
        _gui_log(f"  {n}|  {word_display}  ——  {defs_str}\n")
        _gui_log(f"{'─' * 50}\n")
        _gui_log(f"  ✅ 已保存：{DOC_PATH}\n\n")
        _gui_set_status("待机中")

        try:
            notification.notify(
                title="✅ 笔记已保存",
                message=f"{n}| {word_display} —— {defs_str}"[:200],
                app_name="语言学习助手", timeout=3,
            )
        except Exception:
            pass

    except PermissionError:
        tip = "文档正被 Word 打开，请关闭后再试！"
        _gui_log(f"❌ {tip}\n")
        _gui_set_status("出错")

        try:
            notification.notify(title="⚠️ 保存失败", message=tip, timeout=5)
        except Exception:
            pass

    except Exception as e:
        _gui_log(f"❌ 出错：{e}\n")
        _gui_set_status("出错")

    finally:
        processing_lock.release()
# ============================================================
# GUI
# ============================================================
def start_gui():
    global _gui_log_callback, _gui_count_callback, _gui_status_callback

    root = tk.Tk()
    root.title("PolyglotNote")
    root.geometry("560x480")
    root.minsize(440, 360)
    root.configure(bg="#F0F4F8")

    # ── 顶部标题栏 ───────────────────────────────────────────
    header = tk.Frame(root, bg="#0070C0", height=52)
    header.pack(fill="x")
    header.pack_propagate(False)

    tk.Label(
        header, text="PolyglotNote  语言学习助手",
        font=("等线", 14, "bold"), fg="white", bg="#0070C0",
    ).pack(side="left", padx=16, pady=12)

    count_var = tk.StringVar(value="共 0 条")
    tk.Label(
        header, textvariable=count_var,
        font=("等线", 10), fg="#CCE5FF", bg="#0070C0",
    ).pack(side="right", padx=16, pady=12)

    # ── 文档路径提示 ─────────────────────────────────────────
    info_frame = tk.Frame(root, bg="#E8F0FE", pady=4)
    info_frame.pack(fill="x")

    tk.Label(
        info_frame,
        text=f"📂  {DOC_PATH}",
        font=("等线", 9), fg="#0070C0", bg="#E8F0FE",
        anchor="w",
    ).pack(padx=12, fill="x")

    exe_ok = os.path.exists(ESHELPER_EXE)
    tk.Label(
        info_frame,
        text=("✅  西语助手就绪" if exe_ok else f"⚠️  未找到西语助手：{ESHELPER_EXE}"),
        font=("等线", 9),
        fg=("#2E7D32" if exe_ok else "#B71C1C"),
        bg="#E8F0FE", anchor="w",
    ).pack(padx=12, fill="x", pady=(0, 4))

    # ── 日志区 ───────────────────────────────────────────────
    log_frame = tk.Frame(root, bg="#F0F4F8")
    log_frame.pack(fill="both", expand=True, padx=10, pady=(8, 4))

    log_box = scrolledtext.ScrolledText(
        log_frame,
        font=("等线", 10),
        bg="#FAFAFA", fg="#1A1A2E",
        relief="flat", bd=1,
        wrap="word",
        state="normal",
    )
    log_box.pack(fill="both", expand=True)

    # ── 底部：状态 + 按钮 ────────────────────────────────────
    bottom = tk.Frame(root, bg="#F0F4F8")
    bottom.pack(fill="x", padx=10, pady=(4, 8))

    status_var = tk.StringVar(value="待机中")
    tk.Label(
        bottom, textvariable=status_var,
        font=("等线", 9), fg="#607D8B", bg="#F0F4F8", anchor="w",
    ).pack(side="left")

    def clear_log():
        log_box.delete("1.0", tk.END)

    tk.Button(
        bottom, text="清空日志",
        font=("等线", 9), relief="flat",
        bg="#E3EAF5", fg="#0070C0", cursor="hand2",
        command=clear_log,
    ).pack(side="right", padx=(6, 0))

    def on_closing():
        _gui_log("正在保存时间戳并安全退出…\n")
        append_timestamp()
        # ── FIX: 退出前清理所有键盘钩子 ──────────────────────
        try:
            keyboard.unhook_all()
        except Exception:
            pass
        root.after(300, lambda: (root.destroy(), os._exit(0)))

    tk.Button(
        bottom, text="安全退出",
        font=("等线", 9, "bold"), relief="flat",
        bg="#FFEBEE", fg="#C62828", cursor="hand2",
        command=on_closing,
    ).pack(side="right")

    root.protocol("WM_DELETE_WINDOW", on_closing)
    # ── FIX: ESC 只绑定窗口焦点，不全局监听 ──────────────────
    root.bind("<Escape>", lambda _e: on_closing())

    # ── GUI 回调注入 ─────────────────────────────────────────
    def _log(text: str):
        # ── FIX: str() 保护，防止非字符串类型导致 crash ──────
        root.after(0, lambda t=str(text): (
            log_box.insert(tk.END, t),
            log_box.see(tk.END),
        ))

    def _set_count(n: int):
        root.after(0, lambda: count_var.set(f"共 {n} 条"))

    def _set_status(msg: str):
        root.after(0, lambda: status_var.set(msg))

    _gui_log_callback    = _log
    _gui_count_callback  = _set_count
    _gui_status_callback = _set_status

    # 初始日志
    _log("─" * 50 + "\n")
    _log(" S+D+C   西语助手查词 + Google多词性笔记\n")
    _log(" S+C     仅 Google 多词性西班牙语笔记\n")
    _log(" E+C     仅 Google 多词性英语笔记\n")
    _log("─" * 50 + "\n\n")

    # ── 注册热键（不含 ESC）──────────────────────────────────
    keyboard.add_hotkey("s+d+c", lambda: process("es", open_app=True))
    keyboard.add_hotkey("s+c",   lambda: process("es", open_app=False))
    keyboard.add_hotkey("e+c",   lambda: process("en", open_app=False))

    root.mainloop()


if __name__ == "__main__":
    start_gui()
